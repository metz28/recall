"""
Tests for document loading service
"""
import pytest
import tempfile
import os
from pathlib import Path
from docx import Document
import fitz  # PyMuPDF

from services.document_loader import (
    load_pdf,
    load_docx,
    load_text,
    load_document
)


class TestLoadPDF:
    """Test PDF loading functionality"""

    def test_load_simple_pdf(self):
        """Test loading a simple PDF with text"""
        # Create a temporary PDF
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)  # Close file descriptor

        try:
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Hello, PDF World!")
            doc.save(tmp_path)
            doc.close()

            # Load the PDF
            text = load_pdf(tmp_path)

            assert "Hello, PDF World!" in text
            assert len(text) > 0
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_multipage_pdf(self):
        """Test loading a PDF with multiple pages"""
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)

        try:
            doc = fitz.open()
            page1 = doc.new_page()
            page1.insert_text((72, 72), "Page 1 content")
            page2 = doc.new_page()
            page2.insert_text((72, 72), "Page 2 content")
            doc.save(tmp_path)
            doc.close()

            # Load the PDF
            text = load_pdf(tmp_path)

            assert "Page 1 content" in text
            assert "Page 2 content" in text
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_empty_pdf(self):
        """Test loading an empty PDF"""
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)

        try:
            doc = fitz.open()
            doc.new_page()  # Empty page
            doc.save(tmp_path)
            doc.close()

            # Load the PDF
            text = load_pdf(tmp_path)

            # Should return empty or whitespace string
            assert len(text.strip()) == 0
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_pdf_with_special_characters(self):
        """Test loading PDF with special characters"""
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)

        try:
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "Special: é, ñ")
            doc.save(tmp_path)
            doc.close()

            # Load the PDF
            text = load_pdf(tmp_path)

            # Should contain the text (encoding might vary)
            assert len(text) > 0
        finally:
            Path(tmp_path).unlink(missing_ok=True)


class TestLoadDOCX:
    """Test DOCX loading functionality"""

    def test_load_simple_docx(self):
        """Test loading a simple DOCX with text"""
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        try:
            doc = Document()
            doc.add_paragraph("Hello, DOCX World!")
            doc.save(tmp_path)

            # Load the DOCX
            text = load_docx(tmp_path)

            assert "Hello, DOCX World!" in text
            assert len(text) > 0
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_multiparagraph_docx(self):
        """Test loading DOCX with multiple paragraphs"""
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        try:
            doc = Document()
            doc.add_paragraph("Paragraph 1")
            doc.add_paragraph("Paragraph 2")
            doc.add_paragraph("Paragraph 3")
            doc.save(tmp_path)

            # Load the DOCX
            text = load_docx(tmp_path)

            assert "Paragraph 1" in text
            assert "Paragraph 2" in text
            assert "Paragraph 3" in text
            # Paragraphs should be separated by newlines
            assert "\n" in text
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_empty_docx(self):
        """Test loading an empty DOCX"""
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        try:
            doc = Document()
            doc.save(tmp_path)

            # Load the DOCX
            text = load_docx(tmp_path)

            # Should return empty string
            assert len(text.strip()) == 0
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_docx_with_formatting(self):
        """Test loading DOCX with various formatting (bold, italic, etc.)"""
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        try:
            doc = Document()
            para = doc.add_paragraph()
            para.add_run("Bold text").bold = True
            para.add_run(" and ")
            para.add_run("italic text").italic = True
            doc.save(tmp_path)

            # Load the DOCX
            text = load_docx(tmp_path)

            # Formatting should be stripped, but text should remain
            assert "Bold text" in text
            assert "italic text" in text
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_docx_with_special_characters(self):
        """Test loading DOCX with special characters"""
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        try:
            doc = Document()
            doc.add_paragraph("Special: é, ñ")
            doc.save(tmp_path)

            # Load the DOCX
            text = load_docx(tmp_path)

            assert "Special:" in text
            assert len(text) > 0
        finally:
            Path(tmp_path).unlink(missing_ok=True)


class TestLoadText:
    """Test plain text loading functionality"""

    def test_load_simple_text(self):
        """Test loading a simple text file"""
        fd, tmp_path = tempfile.mkstemp(suffix=".txt")

        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write("Hello, text world!")

            # Load the text
            text = load_text(tmp_path)

            assert text == "Hello, text world!"
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_multiline_text(self):
        """Test loading text with multiple lines"""
        fd, tmp_path = tempfile.mkstemp(suffix=".txt")

        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write("Line 1\nLine 2\nLine 3")

            # Load the text
            text = load_text(tmp_path)

            assert "Line 1" in text
            assert "Line 2" in text
            assert "Line 3" in text
            assert text.count("\n") == 2
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_empty_text(self):
        """Test loading an empty text file"""
        fd, tmp_path = tempfile.mkstemp(suffix=".txt")
        os.close(fd)

        try:
            # Load the text
            text = load_text(tmp_path)

            assert text == ""
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_text_with_special_characters(self):
        """Test loading text with special characters"""
        fd, tmp_path = tempfile.mkstemp(suffix=".txt")

        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write("Special: €, ©, ™, é, ñ, 中文")

            # Load the text
            text = load_text(tmp_path)

            assert "Special:" in text
            assert "€" in text
            assert "中文" in text
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_text_large_file(self):
        """Test loading a large text file"""
        fd, tmp_path = tempfile.mkstemp(suffix=".txt")

        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                # Write 1000 lines
                for i in range(1000):
                    f.write(f"Line {i}\n")

            # Load the text
            text = load_text(tmp_path)

            assert "Line 0" in text
            assert "Line 999" in text
            assert text.count("\n") == 1000
        finally:
            Path(tmp_path).unlink(missing_ok=True)


class TestLoadDocument:
    """Test the main load_document function"""

    def test_load_document_pdf(self):
        """Test loading a PDF through load_document"""
        fd, tmp_path = tempfile.mkstemp(suffix=".pdf")
        os.close(fd)

        try:
            doc = fitz.open()
            page = doc.new_page()
            page.insert_text((72, 72), "PDF content")
            doc.save(tmp_path)
            doc.close()

            # Load through main function
            content, file_type = load_document(tmp_path)

            assert "PDF content" in content
            assert file_type == "pdf"
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_document_docx(self):
        """Test loading a DOCX through load_document"""
        fd, tmp_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)

        try:
            doc = Document()
            doc.add_paragraph("DOCX content")
            doc.save(tmp_path)

            # Load through main function
            content, file_type = load_document(tmp_path)

            assert "DOCX content" in content
            assert file_type == "docx"
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_document_txt(self):
        """Test loading a TXT through load_document"""
        fd, tmp_path = tempfile.mkstemp(suffix=".txt")

        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write("Text content")

            # Load through main function
            content, file_type = load_document(tmp_path)

            assert content == "Text content"
            assert file_type == "txt"
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_document_md(self):
        """Test loading a Markdown file through load_document"""
        fd, tmp_path = tempfile.mkstemp(suffix=".md")

        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write("# Markdown content\n\nWith some text.")

            # Load through main function
            content, file_type = load_document(tmp_path)

            assert "# Markdown content" in content
            assert file_type == "md"
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_document_unsupported_type(self):
        """Test loading an unsupported file type"""
        fd, tmp_path = tempfile.mkstemp(suffix=".xyz")
        os.close(fd)

        try:
            # Should raise ValueError
            with pytest.raises(ValueError, match="Unsupported file type"):
                load_document(tmp_path)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

    def test_load_document_case_insensitive(self):
        """Test that file extension matching is case-insensitive"""
        fd, tmp_path = tempfile.mkstemp(suffix=".TXT")

        try:
            with os.fdopen(fd, "w", encoding="utf-8") as f:
                f.write("Uppercase extension")

            # Load through main function (extension is .TXT)
            content, file_type = load_document(tmp_path)

            assert content == "Uppercase extension"
            assert file_type == "txt"
        finally:
            Path(tmp_path).unlink(missing_ok=True)
